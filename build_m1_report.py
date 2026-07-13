# -*- coding: utf-8 -*-
"""Build the M1 final report as both Markdown (working draft) and Word (.docx).
Single source of truth -> renders to docs/M1-FINAL-REPORT.md and M1-FINAL-REPORT.docx.
Run: python -X utf8 build_m1_report.py
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPORT_DATE = "2026-07-13"
TITLE = "Family Meal Log MVP"
SUBTITLE = "Milestone M1 — Identity & Family Profiles（含 M1.1 加固）"
BLUE_DARK = RGBColor(0x1A, 0x3A, 0x6B)
BLUE_HEAD = RGBColor(0x2B, 0x5C, 0x8F)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
FONT_NAME = "宋体"
FONT_SIZE = 12

# ---------------------------------------------------------------------------
# Content model: list of sections; each section is a list of blocks.
# block kinds: ('h2', text) | ('p', text) | ('bullet', text)
#             | ('table', [headers], [[row], ...])
#             | ('code', text)
# ---------------------------------------------------------------------------

SECTIONS = [
    ("1. M1 完成总结", [
        ("p", "M1 在已验收的 M0 基础之上实现，目标是为「家庭用餐记录」小程序建立单一所有者（single-owner）身份模型与家庭成员档案（family profiles）。本里程碑不涉及多人共享、邀请或角色权限，且档案删除（归档/软删除）明确延后至后续里程碑。"),
        ("p", "M1 交付的核心成果："),
        ("bullet", "新增独立云函数 profileApi（list / get / create / update / setDefault），不复用 mealApi；login 与 profileApi 均通过共享运行时（shared runtime）派生身份与档案逻辑。"),
        ("bullet", "服务端可信身份：openid 仅由 cloud.getWXContext().OPENID 在服务端取得，客户端永不发送、接收或存储 openid；所有对外 DTO 不含 openid / ownerOpenid。"),
        ("bullet", "共享运行时确定性打包：shared/*.ts 编译为 CommonJS 后复制进各云函数 lib/shared/，无符号链接、Windows 兼容、生成物不入库。"),
        ("bullet", "客户端身份/档案 UI：首次引导（onboarding）、档案列表/创建/编辑/选择/设默认，首页展示当前活动档案。"),
        ("bullet", "自动化验收：npm run validate 共 91 项检查全部通过（M0 基线 67 + M1/M1.1 共 24）。其中「跨用户隔离」「共享运行时打包」「默认档案单一事实来源」「best-effort 请求幂等（尽力式请求幂等）」四组关键验收均通过。"),
        ("p", "M1.1 加固（合入 main 前的窄范围修正，详见第 16 节）：(1) 移除基于姓名的档案去重，改为客户端 requestId 的 best-effort 请求幂等（尽力式请求幂等：客户端 in-flight 防护 + 服务端请求重放处理）+ UI 防抖；(2) 默认档案单一事实来源——仅以 users.defaultFamilyProfileId 持久化，family_profiles 不存 isDefault，对外 isDefault 由 DTO 计算。"),
        ("p", "结论：依据自动化验收门槛（owner-isolation、shared-runtime packaging、default-profile single-source-of-truth、best-effort request idempotency 必须全部通过），M1 + M1.1 关键验收达成。**代码与文档已就绪，可进入产品负责人的 DevTools / 真机手动验收（见第 12、17、18 节）**；当前自动化环境无法代跑手动/真机测试，故不宣称产品已验收通过。"),
        ("p", "附带一项仓库卫生修正：.gitignore 现忽略 .workbuddy/（项目记忆目录，不应提交）。"),
    ]),
    ("2. 变更区域仓库目录树", [
        ("p", "以下为 M1 实际变更/新增的目录与文件（省略未变动项）："),
        ("code",
"""Family Meal Log MVP/
├─ .gitignore                      # 修正：仅忽略生成物 lib/shared/，提交云函数源码 .js
├─ package.json                    # 新增 build:shared / clean 脚本
├─ README.md                       # M0/M1 里程碑表、仓库结构、校验说明
├─ shared/                         # ★ 共享运行时（单一事实来源）
│  ├─ index.ts
│  ├─ repository.ts                # Repository 接口 + ServiceError
│  ├─ repository-memory.ts         # InMemoryRepository（测试用）
│  ├─ tsconfig.json
│  ├─ dist/                        # 编译产出（git 忽略）
│  └─ services/
│     ├─ user-service.ts           # upsertUser（幂等登录）
│     ├─ profile-service.ts        # 档案校验/创建/更新/设默认/列表
│     └─ session.ts                # resolveActiveProfile 纯函数
├─ cloudfunctions/
│  ├─ login/                       # 改用共享运行时；去除响应中的 openid
│  │  ├─ index.js  cloudbase-repository.js  package.json  config.json
│  ├─ profileApi/                  # ★ 新增云函数
│  │  ├─ index.js  cloudbase-repository.js  package.json  config.json
│  ├─ mealApi/                     # 修订：响应去除 openid
│  └─ aiAnalyze/                   # 随 gitignore 修正一并纳入版本控制（M0 遗留）
├─ miniprogram/
│  ├─ app.ts  app.json             # onLaunch 初始化云 + 可选 loadSession
│  ├─ typings/index.d.ts           # 移除 openid；新增 ClientFamilyProfile / FamilyRelation
│  ├─ config/labels.ts             # ★ RELATION_LABELS / RELATION_OPTIONS（中文）
│  ├─ services/
│  │  ├─ auth.ts                   # ★ login() 封装
│  │  ├─ profile.ts                # ★ 档案接口封装
│  │  └─ session.ts                # ★ 活动档案解析与本地存储
│  └─ pages/
│     ├─ home/                     # 展示活动档案 / 引导 / 离线；保留记餐占位
│     ├─ profiles/                 # ★ 档案列表/选择/设默认（防重复点击）
│     └─ profile-edit/             # ★ 创建/编辑表单（姓名 + 关系）
├─ scripts/
│  ├─ build-shared.mjs             # ★ 共享运行时编译+复制
│  └─ validate.mjs                 # 扩展 16 项 M1 验收
└─ docs/
   ├─ SECURITY.md                  # ★ 新增：信任模型/集合/索引/安全规则
   ├─ MANUAL_TEST_CHECKLIST.md     # ★ 新增：DevTools + 真机清单
   ├─ ARCHITECTURE.md              # 更新 M1 状态/页面/服务/打包/安全/决策
   ├─ DATA_MODEL.md                # 更新索引/校验规则/删除延后/本地日
   ├─ USER_FLOWS.md                # 修正 openid 表述，补充 M1 流程
   ├─ DEVELOPMENT_PLAN.md          # M1 标记 ✅
   └─ PRODUCT_REQUIREMENTS.md      # 补充排除项（删除延后/营养目标延后）
"""),
    ]),
    ("3. 创建与修改文件完整清单", [
        ("p", "M1 最初影响 **47 个文件**（31 新增、16 修改）。M1.1 在此基础上追加修改了 **18 个文件**（全部为 M1 既有文件，即在相同代码/文档之上叠加的加固层）。本次最终 closeout（本报告交付）又修改了一组文档/报告文件（见 §3.6）。下方按类别列示实质性文件；逐文件完整清单以 `git diff --stat` 为准。"),
        ("h2", "3.1 共享运行时（新增）"),
        ("table", ["文件", "状态", "用途"], [
            ["shared/repository.ts", "新增", "Repository 接口、ServiceError、错误码"],
            ["shared/repository-memory.ts", "新增", "内存实现，供自动化测试使用"],
            ["shared/index.ts", "修改", "导出 repository / service 模块"],
            ["shared/services/user-service.ts", "新增", "upsertUser 幂等登录"],
            ["shared/services/profile-service.ts", "新增", "档案校验、CRUD、设默认、列表面向客户端转换"],
            ["shared/services/session.ts", "新增", "resolveActiveProfile 纯函数（优先级逻辑）"],
        ]),
        ("h2", "3.2 云函数"),
        ("table", ["文件", "状态", "用途"], [
            ["cloudfunctions/login/index.js", "修改", "改用共享运行时；响应去除 openid"],
            ["cloudfunctions/login/cloudbase-repository.js", "新增", "基于 wx-server-sdk 的 Repository 实现"],
            ["cloudfunctions/profileApi/index.js", "新增", "profileApi 分发 list/get/create/update/setDefault"],
            ["cloudfunctions/profileApi/cloudbase-repository.js", "新增", "同 login 的仓储实现"],
            ["cloudfunctions/profileApi/package.json", "新增", "依赖 wx-server-sdk"],
            ["cloudfunctions/profileApi/config.json", "新增", "云函数配置"],
            ["cloudfunctions/mealApi/index.js", "修改", "响应去除 openid"],
            ["cloudfunctions/aiAnalyze/index.js", "新增(纳入)", "M0 遗留源码随 gitignore 修正入库"],
        ]),
        ("h2", "3.3 客户端"),
        ("table", ["文件", "状态", "用途"], [
            ["miniprogram/app.ts", "修改", "onLaunch 初始化云 + 可选 loadSession"],
            ["miniprogram/app.json", "修改", "注册 profiles / profile-edit 页面"],
            ["miniprogram/typings/index.d.ts", "修改", "移除 openid；新增 ClientFamilyProfile/FamilyRelation"],
            ["miniprogram/config/labels.ts", "新增", "关系中文标签与选择器选项"],
            ["miniprogram/services/auth.ts", "新增", "login() 封装"],
            ["miniprogram/services/profile.ts", "新增", "档案接口封装 + ApiResult"],
            ["miniprogram/services/session.ts", "新增", "活动档案解析/本地存储/loadSession"],
            ["miniprogram/pages/home/home.ts/.wxml/.wxss", "修改", "活动档案/引导/离线；记餐占位"],
            ["miniprogram/pages/profiles/*", "新增", "档案列表/选择/设默认"],
            ["miniprogram/pages/profile-edit/*", "新增", "创建/编辑表单"],
        ]),
        ("h2", "3.4 构建 / 校验 / 文档 / 配置"),
        ("table", ["文件", "状态", "用途"], [
            ["scripts/build-shared.mjs", "新增", "共享运行时编译+复制打包"],
            ["scripts/validate.mjs", "修改", "扩展 16 项 M1 验收"],
            ["package.json", "修改", "新增 build:shared / clean"],
            [".gitignore", "修改", "仅忽略生成物；忽略 .workbuddy/"],
            ["README.md", "修改", "里程碑表/结构/校验说明"],
            ["docs/SECURITY.md", "新增", "信任模型/集合/索引/安全规则"],
            ["docs/MANUAL_TEST_CHECKLIST.md", "新增", "DevTools + 真机清单"],
            ["docs/ARCHITECTURE.md", "修改", "M1 状态/页面/服务/打包/安全/决策"],
            ["docs/DATA_MODEL.md", "修改", "索引/校验规则/删除延后/本地日"],
            ["docs/USER_FLOWS.md", "修改", "修正 openid 表述，补充 M1 流程"],
            ["docs/DEVELOPMENT_PLAN.md", "修改", "M1 标记 ✅"],
            ["docs/PRODUCT_REQUIREMENTS.md", "修改", "补充排除项"],
        ]),
        ("h2", "3.5 M1.1 修改文件（18 个，均叠加于 M1 既有文件）"),
        ("table", ["文件", "状态", "M1.1 变更"], [
            ["shared/types.ts", "修改", "新增 IdempotencyKey 接口；注明 FamilyProfile 不存 isDefault"],
            ["shared/repository.ts", "修改", "Repository 增加 findIdempotencyKey / saveIdempotencyKey"],
            ["shared/repository-memory.ts", "修改", "内存实现幂等键存取"],
            ["shared/services/profile-service.ts", "修改", "移除姓名去重；createProfile 支持 requestId 幂等；DTO 计算 isDefault"],
            ["cloudfunctions/profileApi/cloudbase-repository.js", "修改", "idempotency_keys 集合的查/写实现"],
            ["cloudfunctions/profileApi/index.js", "修改", "currentDefaultId 助手；透传 requestId"],
            ["cloudfunctions/login/cloudbase-repository.js", "修改", "接口对齐补齐幂等键方法"],
            ["typings/index.d.ts", "修改", "ClientFamilyProfile 增 isDefault（服务端计算）"],
            ["miniprogram/services/profile.ts", "修改", "newRequestId()；createProfile 传 requestId"],
            ["miniprogram/pages/profile-edit/profile-edit.ts", "修改", "onLoad 生成 createRequestId；onSubmit 透传"],
            ["scripts/validate.mjs", "修改", "14a/14b/14c + 15a–15f 替换姓名去重测试；打包/卫生重编号"],
            ["docs/DATA_MODEL.md", "修改", "idempotency_keys、默认 SSoT、移除姓名去重"],
            ["docs/ARCHITECTURE.md", "修改", "§5b 幂等与默认 SSoT、§12 M1.1 决策、集合/索引"],
            ["docs/USER_FLOWS.md", "修改", "创建流程 requestId、允许重名、默认计算与回退"],
            ["docs/SECURITY.md", "修改", "集合/索引、创建幂等与默认 SSoT 说明"],
            ["build_m1_report.py", "修改", "并入 M1.1 修正、测试状态与 git 复核"],
            ["docs/M1-FINAL-REPORT.md", "修改", "重新生成（Markdown 草稿）"],
            ["M1-FINAL-REPORT.docx", "修改", "重新生成（Word，蓝色主题）"],
        ]),
        ("h2", "3.6 最终 closeout 修改文件（本报告交付，M1 final closeout）"),
        ("table", ["文件", "状态", "closeout 变更"], [
            ["docs/DATA_MODEL.md", "修改", "幂等表述改为 best-effort（尽力式请求幂等）；索引改 composite 非唯一"],
            ["docs/ARCHITECTURE.md", "修改", "§5b 增加 best-effort 术语；索引标注 currently non-unique"],
            ["docs/SECURITY.md", "修改", "§3 明确 best-effort 请求幂等（非原子）"],
            ["docs/USER_FLOWS.md", "修改", "创建流程增加 best-effort 术语"],
            ["docs/DEVELOPMENT_PLAN.md", "修改", "M2 范围对齐「食物目录与份量单位」+ 排除项；M1 验收措辞"],
            ["README.md", "修改", "里程碑表 M2/M3 措辞 + M2 排除说明"],
            ["docs/MANUAL_TEST_RESULTS.md", "新增", "手动测试结果表（14 项全 Pending）"],
            ["build_m1_report.py", "修改", "报告源：幂等表述/路线图/Git 复核/文件计数"],
            ["docs/M1-FINAL-REPORT.md", "重新生成", "本报告 Markdown 草稿"],
            ["M1-FINAL-REPORT.docx", "重新生成", "本报告 Word（蓝色主题）"],
        ]),
    ]),
    ("4. 身份与数据模型", [
        ("h2", "4.1 users 集合"),
        ("table", ["字段", "类型", "说明"], [
            ["_id", "string", "云数据库文档 ID（对外仅返回此内部 ID）"],
            ["openid", "string", "微信身份；唯一索引；仅服务端写入，客户端不可见"],
            ["unionid", "string?", "跨应用 UnionID（可选）"],
            ["defaultFamilyProfileId", "string?", "当前默认档案 ID；登录/设默认时维护"],
            ["createdAt / updatedAt", "number", "服务端时间戳（毫秒）"],
        ]),
        ("p", "登录幂等：以 openid 为键 upsert，首次 isNew=true 并写入时间戳；并发竞态下由唯一索引保障单文档（重复写入以相同 openid 命中既有文档）。"),
        ("h2", "4.2 family_profiles 集合"),
        ("table", ["字段", "类型", "说明"], [
            ["_id", "string", "档案文档 ID"],
            ["ownerOpenid", "string", "所有者 openid；仅服务端设置，客户端不可伪造"],
            ["name", "string", "姓名；trim 后非空、≤30 字符；不作唯一键（同一所有者可重名）"],
            ["relation", "enum", "self / spouse / child / parent / other"],
            ["createdAt / updatedAt", "number", "服务端时间戳"],
        ]),
        ("p", "校验规则（normalizeProfileInput）：姓名去除首尾空格、限长 30、非空；relation 必须为白名单枚举；除 name/relation 外的字段一律丢弃，ownerOpenid 由服务端写入。更新时仅 name/relation 可改，其余（尤其 ownerOpenid）不可变。档案删除不在 M1 范围，留待后续以归档/软删除实现。"),
        ("p", "【M1.1】family_profiles 不再存储 isDefault 字段；对外 DTO 的 isDefault 由服务端按 profile.id === users.defaultFamilyProfileId 计算得出（详见第 16 节）。姓名不作唯一键，不做基于姓名的去重，同一所有者可创建重名档案。"),
        ("h2", "4.2b idempotency_keys 集合（M1.1）"),
        ("table", ["字段", "类型", "说明"], [
            ["_id", "string", "记录文档 ID"],
            ["ownerOpenid", "string", "所有者 openid；仅服务端写入"],
            ["operation", "string", "逻辑操作，如 create（家庭档案）"],
            ["requestId", "string", "客户端生成的高熵 ID，一次编辑会话内稳定"],
            ["resultId", "string", "首个请求所创建实体的 _id"],
            ["createdAt", "number", "服务端时间戳"],
        ]),
        ("p", "以三元组 (ownerOpenid, operation, requestId) 标识一次逻辑意图：相同 requestId 的重复请求通常返回原档案（resultId），不新建；不同 requestId 视为新意图（即便姓名/关系相同）。requestId 以可信 ownerOpenid 作用域隔离，不同用户复用同一字符串不冲突。这是 **best-effort 请求幂等（尽力式请求幂等）**——客户端 in-flight 防护 + 服务端请求重放处理——**不是**严格/原子保证：read→create→write-key 非原子，存在残留并发竞态（详见第 16 节与 docs）。"),
        ("p", "本地自然日约定：用餐记录（M2）按 YYYY-MM-DD 00:00–23:59（所有者所在时区）分区存储，该约定已在 DATA_MODEL 登记，M1 仅记录规则不落地数据。"),
        ("h2", "4.3 索引"),
        ("table", ["集合", "字段", "类型", "用途"], [
            ["users", "openid", "唯一", "幂等登录、按身份查用户"],
            ["family_profiles", "ownerOpenid + createdAt", "普通", "按所有者列出/隔离档案，确定性排序"],
            ["idempotency_keys", "ownerOpenid + operation + requestId", "复合（当前非唯一；M3 升级为唯一索引作原子闸）", "创建操作的 best-effort 请求幂等"],
        ]),
    ]),
    ("5. 云函数契约", [
        ("h2", "5.1 login"),
        ("bullet", "从 cloud.getWXContext().OPENID 取身份（绝不信任客户端传入）。"),
        ("bullet", "调用 upsertUser(repo, openid, unionid?) → { user, isNew }。"),
        ("bullet", "响应：{ ok:true, user:{ id, defaultFamilyProfileId }, isNewUser }。不包含 openid。"),
        ("h2", "5.2 profileApi 动作"),
        ("table", ["动作", "说明", "对外字段（toClientProfile）"], [
            ["list", "按 OPENID 列出本人档案，按 createdAt 升序；isDefault 由 DTO 计算", "id, name, relation, isDefault(计算), createdAt"],
            ["get", "取单条，校验归属；isDefault 由 DTO 计算", "同上"],
            ["create", "规范化输入；服务端写 ownerOpenid；首个档案自动默认；按 requestId 的 best-effort 请求幂等（尽力式请求幂等）；不按姓名去重", "返回新建（或幂等命中的原）档案"],
            ["update", "校验归属；仅可改 name/relation", "返回更新后档案"],
            ["setDefault", "校验归属；仅更新 users.defaultFamilyProfileId，不改任何档案文档", "返回目标档案（isDefault 计算为 true）"],
        ]),
        ("p", "【M1.1】create 携带 event.requestId，服务端以 (ownerOpenid, 'create', requestId) 查 idempotency_keys：命中则返回原档案，未命中则创建并记录键；不再进行基于姓名的去重。这是 **best-effort 请求幂等**：read→create→write-key 非原子，理论上并发同 requestId 可能各建一条，但 UI in-flight 防护 + 单用户点击并发已将其缓解（残留竞态见第 16 节、docs）。isDefault 一律由 DTO 计算（profile.id === user.defaultFamilyProfileId），服务端不持久化该字段。"),
        ("p", "错误映射：ServiceError（not_found / forbidden / validation / invalid_input）→ { ok:false, error:{ code, message } }。mealApi 同步去除响应中的 openid（M0 遗留修正）。"),
    ]),
    ("6. 共享运行时打包", [
        ("p", "目标：shared/*.ts 为单一事实来源，编译后确定性地复制进各云函数，避免符号链接、兼容 Windows、生成物不入库。"),
        ("bullet", "编译：tsc -p shared/tsconfig.json → shared/dist（CommonJS）。"),
        ("bullet", "复制：scripts/build-shared.mjs 将 shared/dist 递归复制到 cloudfunctions/<fn>/lib/shared/，并为该函数写入 { type:'commonjs' } package.json；覆盖 login 与 profileApi（亦可供其它函数复用）。"),
        ("bullet", "git 忽略：cloudfunctions/*/lib/shared/ 与 shared/dist/ 均忽略；仅提交云函数源码 .js（修正 M0 的 cloudfunctions/**/*.js 误忽略）。"),
        ("bullet", "校验：validate 的 15a–d 检查 login/profileApi 是否打包共享运行时且确实引用；另含 gitignore 与生成物忽略检查。"),
    ]),
    ("7. 客户端架构", [
        ("h2", "7.1 服务层"),
        ("bullet", "services/auth.ts：login() 封装 login 云函数，返回 { userId, defaultFamilyProfileId, isNewUser }。"),
        ("bullet", "services/profile.ts：listProfiles / createProfile / updateProfile / setDefaultProfile / getProfile 封装 profileApi；ApiResult 非泛型；映射 invalid_input。"),
        ("bullet", "services/session.ts：getActiveProfileId / setActiveProfileId（wx.storage）；客户端 resolveActiveProfile；loadSession(app)；selectActiveProfile。"),
        ("bullet", "config/labels.ts：RELATION_LABELS（中文展示）与 RELATION_OPTIONS（选择器）。"),
        ("h2", "7.2 页面"),
        ("bullet", "pages/home：展示当前活动档案 / 引导 / 离线横幅；保留「记餐」占位入口。"),
        ("bullet", "pages/profiles：档案列表、选择、设默认（防重复点击）。"),
        ("bullet", "pages/profile-edit：创建/编辑表单（姓名输入 + 关系选择器），提交防重。"),
        ("p", "app.ts 在 onLaunch 初始化云环境，并可选地 loadSession 预填全局活动档案。typings/index.d.ts 已移除 openid，新增 ClientFamilyProfile 与 FamilyRelation 全局类型。"),
    ]),
    ("8. 活动档案解析与引导流程", [
        ("p", "解析优先级（共享纯函数 resolveActiveProfile，客户端镜像同一逻辑）："),
        ("bullet", "1) 本地记住的档案 ID —— 若仍在列表内则采用；"),
        ("bullet", "2) 服务端 defaultFamilyProfileId；"),
        ("bullet", "3) 列表首个档案；"),
        ("bullet", "4) 无档案 → 进入引导（onboarding）。"),
        ("p", "本地仅存储档案 ID（wx.storage），不缓存 openid 或敏感字段。首次运行（无档案）进入引导创建首个档案，该档案自动设为默认。session.loadSession 在启动后将活动档案写入 app.globalData。"),
    ]),
    ("9. 安全模型", [
        ("bullet", "服务端可信身份：openid 仅来自 cloud.getWXContext().OPENID；客户端永不发送/接收/存储 openid。"),
        ("bullet", "客户端安全 DTO：对外结构排除 openid 与 ownerOpenid，仅返回内部文档 ID 与档案摘要。"),
        ("bullet", "跨用户隔离：list / get / update / setDefault 全部按 OPENID 校验归属，越权返回 forbidden / not_found（ServiceError）。"),
        ("bullet", "CloudBase 安全规则：相关集合 read:false、write:false，所有访问经云函数；users.openid 与 family_profiles.ownerOpenid 建索引以支持隔离查询。"),
        ("bullet", "无密钥入库：校验含 no-secrets 检查；SECURITY.md 提供开发/生产一致性与卫生清单。"),
    ]),
    ("10. 状态与错误处理矩阵", [
        ("table", ["场景", "行为"], [
            ["云环境未配置", "首页显示离线横幅；服务失败优雅降级，不崩溃"],
            ["登录失败", "提示错误，不加载档案"],
            ["列表失败", "显示重试/空态"],
            ["创建/更新失败", "Toast 报错，保留表单内容"],
            ["重复点击提交", "客户端按钮防重（in-flight 标记）；服务端按 requestId 的 best-effort 请求幂等（尽力式请求幂等，非原子保证），在通常的单用户操作和网络重试场景下避免重复创建；极端并发情况下仍存在残留竞态。"],
            ["同名档案（不同意图）", "允许创建（不同 requestId）；姓名非唯一键"],
            ["姓名为空/纯空格", "客户端校验 + 服务端 rejection（validation）"],
            ["关系非法", "选择器限定枚举；服务端拒绝"],
            ["本地 ID 失效（档案已删）", "resolveActiveProfile 回退至服务端默认 → 首个 → 引导"],
            ["默认 ID 失效/悬挂", "无档案被标记为默认（安全回退，无需数据修复）"],
            ["操作他人档案", "服务端按 OPENID 校验返回 forbidden（客户端无法伪造归属）"],
            ["无任何档案", "进入引导创建首个档案"],
        ]),
    ]),
    ("11. 自动化校验结果", [
        ("p", "npm run validate 串联 typecheck 与 91 项检查：M0 基线 67 + M1/M1.1 共 24，0 失败。已在最终 closeout 提交前重新执行并通过。分组如下（14a–15f 为 M1.1 新增/调整）："),
        ("table", ["编号", "检查", "分组"], [
            ["1", "login upsert 幂等（同身份仅首次 isNew）", "身份"],
            ["2", "首个档案自动设为默认", "身份"],
            ["3", "用户可创建 ≥2 个档案", "身份"],
            ["4", "档案姓名被 trim", "校验"],
            ["5", "空/纯空格姓名被拒", "校验"],
            ["6", "非法 relation 被拒", "校验"],
            ["7", "未知字段不持久化", "校验"],
            ["8", "客户端传入 ownerOpenid 被忽略", "校验"],
            ["9", "用户不能列出他人档案", "隔离"],
            ["10", "用户不能更新他人档案", "隔离"],
            ["11", "用户不能将他人档案设为默认", "隔离"],
            ["12a–c", "活动档案回退（服务端默认 / 首个 / 引导）", "回退"],
            ["13", "默认在重新登录后持久化", "持久化"],
            ["14a", "同名 + 不同 requestId → 两条不同档案", "幂等(M1.1)"],
            ["14b", "相同 requestId → 返回同一档案（不重复）", "幂等(M1.1)"],
            ["14c", "不同用户复用同一 requestId 不冲突", "幂等(M1.1)"],
            ["15a", "family_profiles 不持久化 isDefault", "默认SSoT(M1.1)"],
            ["15b", "DTO 由 users.defaultFamilyProfileId 计算 isDefault", "默认SSoT(M1.1)"],
            ["15c", "返回列表中恰有一个默认", "默认SSoT(M1.1)"],
            ["15d", "切换默认不改写任何档案文档", "默认SSoT(M1.1)"],
            ["15e", "设新默认改变计算出的 isDefault", "默认SSoT(M1.1)"],
            ["15f", "悬挂/缺失默认 ID → 无档案标记默认（安全回退）", "默认SSoT(M1.1)"],
            ["16a–d", "共享运行时已打包进 login/profileApi 且被引用", "打包"],
            ["—", "gitignore 忽略生成运行时；生成物确被忽略", "打包"],
            ["17", "客户端 globalData / app.ts / login 响应均不含 openid", "卫生"],
        ]),
        ("p", "结论：跨用户隔离（9–11）、best-effort 请求重放处理（14a–c，即尽力式请求幂等测试）、默认档案单一事实来源（15a–f）、共享运行时打包（16a–d）全部通过，M1 + M1.1 关键验收达成。"),
    ]),
    ("12. 手动测试清单", [
        ("p", "详见 docs/MANUAL_TEST_CHECKLIST.md。测试状态严格区分「自动化 / DevTools 手动 / 真机」，详见第 17 节。要点："),
        ("bullet", "A. DevTools（状态：待产品负责人执行）：首次运行进入引导；创建档案后自动默认；再建第二个；切换默认；编辑；选择活动档案；重新编译后默认仍生效；提交空姓名/非法关系被拒；连续点击提交不重复创建；【M1.1】故意创建两条重名档案应均成功。"),
        ("bullet", "B. 真机（状态：待产品负责人执行）：上述流程在手机端复现；离线时首页横幅正确；本地 ID 失效回退正确。"),
        ("p", "注：单设备上难以实测「跨账户隔离」，该路径由自动化测试（9–11）覆盖，真机仅做冒烟。DevTools 与真机测试需在装有微信开发者工具/真机的环境中由人工执行，当前自动化环境无法代为运行（见第 17 节步骤）。"),
    ]),
    ("13. 文档更新与决策记录", [
        ("h2", "13.1 新增文档"),
        ("bullet", "docs/SECURITY.md：信任模型、集合、索引、安全规则、dev/prod、卫生清单。"),
        ("bullet", "docs/MANUAL_TEST_CHECKLIST.md：DevTools + 真机 11 步清单。"),
        ("h2", "13.2 更新文档"),
        ("bullet", "ARCHITECTURE：M1 状态、新增页面/服务、profileApi 分发、§6 打包、§9 安全、§11 决策。"),
        ("bullet", "DATA_MODEL：索引表与竞态说明、User 客户端契约、FamilyProfile M1 校验规则与删除延后、本地日规则。"),
        ("bullet", "USER_FLOWS：修正「客户端在 globalData 存储 openid」的旧表述（§1.5 不再存储），补充 M1 流程。"),
        ("bullet", "DEVELOPMENT_PLAN：M1 标记 ✅ 并注明测试通过。"),
        ("bullet", "PRODUCT_REQUIREMENTS：排除项补充「档案删除延后」「营养目标延后」。"),
        ("bullet", "README：里程碑表 M0/M1 ✅、仓库结构、校验说明。"),
        ("bullet", "【M1.1】DATA_MODEL / ARCHITECTURE / USER_FLOWS / SECURITY：新增 idempotency_keys 集合与索引、requestId 幂等契约与残留竞态说明及 M3 通用设计、默认档案单一事实来源（isDefault 由 DTO 计算、不持久化）、移除姓名去重（允许重名）。"),
        ("h2", "13.3 已记录决策"),
        ("bullet", "单一所有者模型（无多人共享/邀请/角色）。"),
        ("bullet", "档案删除延后（后续以归档/软删除实现）。"),
        ("bullet", "本地自然日 YYYY-MM-DD 约定（为 M2 用餐记录分区）。"),
        ("bullet", "共享运行时打包：编译+复制、无符号链接、生成物不入库，纳入 npm run validate。"),
        ("bullet", "客户端绝不存储 openid；移除旧文档中相关错误表述。"),
    ]),
    ("14. Git 与范围控制", [
        ("p", "分支：feature/m1-identity-family-profiles，自干净的 M0 基线（cc698c2）切出。以下为 `git log --oneline cc698c2..HEAD` 在本报告重新生成前（即提交本 closeout 报告之前）的真实记录，全部为真实 hash，未编造、未 squash："),
        ("h2", "14.1 M1 提交（6 个）"),
        ("table", ["#", "Hash", "提交信息"], [
            ["1", "3f323f5", "M1: shared runtime (repository, user/profile services, session) + build packaging"],
            ["2", "82c8c41", "M1: profileApi cloud function + login uses shared runtime; drop openid from client responses"],
            ["3", "465c868", "M1: client identity/profile UI + auth/profile/session services"],
            ["4", "b79def5", "M1: documentation (architecture, data model, flows, plan, security, manual tests) + ignore .workbuddy"],
            ["5", "d5b1134", "M1: extend validation with 16 M1 acceptance tests"],
            ["6", "acbf11d", "M1: final report (Markdown draft + Word .docx, blue theme)"],
        ]),
        ("h2", "14.2 M1.1 加固提交（4 个，真实 hash）"),
        ("table", ["#", "Hash", "提交信息"], [
            ["7", "a729826", "M1.1: remove name-based dedup (requestId idempotency) + default single source of truth"],
            ["8", "b6aefc6", "M1.1: validation tests for requestId idempotency & default single source of truth"],
            ["9", "1c9243c", "M1.1: docs (data model, architecture, flows, security) reflect idempotency + default SSoT"],
            ["10", "f95054d", "M1.1: reconcile final report (M1.1 corrections, test status, git log)"],
        ]),
        ("p", "说明：上方共 10 个提交，hash 均真实存在，与仓库 `git log --oneline cc698c2..HEAD`（本报告生成前）完全一致。**本报告（本 closeout 重新生成的版本）所在的提交未列入该日志**——因为报告在生成提交之前已写好，其 hash 彼时尚未产生。即：*The commit containing this regenerated report is not included in the log above because the report was generated before that commit was created.* 合并计数：M1 6 个 + M1.1 4 个 = 10 个（本 closeout 的额外提交另计，不计入上方日志）。"),
        ("p", "范围控制：未包含任何 M2+ 代码（无 food/portion/meal/recipe/photo/AI 实现），未实现档案删除、营养目标或医疗字段，未重构 UI。M1 附带修正：M0 的 cloudfunctions/**/*.js 误忽略导致云函数源码从未入库；M1 将 .gitignore 改为仅忽略 lib/shared/，使 login/mealApi/aiAnalyze 源码得以纳入版本控制。"),
    ]),
    ("15. 里程碑路线图（仅规划，不实现）", [
        ("p", "经 PM 复核，M2 不再建议「用餐记录与营养（含 mealApi CRUD）」。约定序列如下，M2 仅聚焦食物目录与份量单位："),
        ("table", ["里程碑", "范围", "明确排除"], [
            ["M2 — Food Catalog & Portion Units", "小型精选系统食物种子数据集；食物搜索与选择；用户自定义临时食物；通用份量单位 + 食物专属份量单位；份量→克换算；单食物实时营养预览；营养记录的来源与版本元数据", "保存餐食；mealApi 增删改查；每日餐史；食谱；照片上传；AI 识别"],
            ["M3 — Manual Meal Logging", "多食物组合成一餐；餐型与日期；服务端校验与重算；保存并重新加载餐食", "继承 M2 排除项；不含编辑/删除（见 M4）"],
            ["M4 — History, Edit & Delete", "每日餐史浏览；编辑与删除餐食", "—"],
            ["M5 — Saved Foods & Recipes", "收藏食物；简单家庭食谱", "—"],
            ["M6 — Photo Upload", "餐食照片上传至 CloudBase Storage", "—"],
            ["M7 — Mock AI Suggestions", "provider-neutral 适配层 + mock 供应商 + 确认/修正 UX", "真实 AI"],
            ["M8 — Real AI Provider", "同一接口后接入真实供应商；密钥仅存 CloudBase 函数环境变量", "—"],
        ]),
        ("p", "M2 详细验收见 docs/DEVELOPMENT_PLAN.md（已标注 M2 仅食物目录与份量单位，不含任何 meal 持久化）。全部里程碑保持单一所有者模型，openid 仅服务端可见。本 closeout 未实现任何 M2+ 代码。"),
    ]),
    ("16. M1.1 加固修正（合入 main 前）", [
        ("p", "在将 M1 合入 main 前执行的窄范围加固，仅含两处设计修正与配套测试/文档，不引入 M2、不实现档案删除/营养目标/医疗字段、不重构 UI、不接真实 AI。"),
        ("h2", "16.1 修正一：移除基于姓名的档案去重（best-effort 请求幂等）"),
        ("bullet", "问题：早期 M1 可能以 name（或 name + relation）作为唯一/幂等键，导致同一所有者无法创建重名档案（如两个都叫「宝宝」的孩子）。"),
        ("bullet", "修正：createProfile 不再按姓名去重；姓名不作唯一键。防止误重复改由两层实现——(a) 客户端提交 in-flight 防抖；(b) 客户端生成 requestId 的 **best-effort 请求幂等（尽力式请求幂等）**：客户端 in-flight 防护 + 服务端请求重放处理，作用域为 (ownerOpenid, operation, requestId)。"),
        ("bullet", "客户端契约：{ action:'create', requestId:'req_<time36>_<random>', profile:{ name, relation } }。相同 requestId 返回原档案；不同 requestId 即便姓名/关系相同也创建新档案。"),
        ("bullet", "落地：新增 idempotency_keys 集合与 IdempotencyKey 类型；Repository 增加 findIdempotencyKey / saveIdempotencyKey；内存实现与 CloudBase 实现均补齐；profileApi.index.js 透传 event.requestId。"),
        ("bullet", "文档化事实（不宣称严格/原子幂等）：(1) 相同 requestId 通常返回原结果；(2) 不同 requestId 允许重名档案；(3) read→create→write-key 非原子，存在残留并发竞态；(4) UI in-flight 防护降低家庭 MVP 的实际风险；(5) 未来通用原子方案应在实体创建前用唯一约束/原子幂等键申领（唯一复合索引，重复键即视为已处理并返回 resultId）；(6) 该更强实现应在此之前或 M3（餐食创建同样需要幂等）引入。"),
        ("bullet", "残留风险（已记录）：read→create→write-key 非原子，理论上两条并发同 requestId 可能各自创建；已由 UI 防抖 + 单用户点击并发缓解。M3 通用方案：将复合索引升级为唯一索引，以「写键」作为原子闸（重复键报错即视为已处理并返回 resultId）。"),
        ("h2", "16.2 修正二：默认档案单一事实来源"),
        ("bullet", "仅以 users.defaultFamilyProfileId 作为默认状态的唯一持久化来源；family_profiles 不存储 isDefault。"),
        ("bullet", "setDefault 仅更新用户文档，不改写任何档案文档（切换默认无需多文档写入）。"),
        ("bullet", "对外 isDefault 一律在 DTO 计算：toClientProfile(p, defaultProfileId) 以 p.id === defaultProfileId 得出。"),
        ("bullet", "悬挂/缺失默认 ID → 无档案被标记默认（安全回退，无需数据修复）。"),
        ("h2", "16.3 变更文件清单（M1.1）"),
        ("table", ["文件", "状态", "M1.1 变更"], [
            ["shared/types.ts", "修改", "新增 IdempotencyKey 接口；注明 FamilyProfile 不存 isDefault"],
            ["shared/repository.ts", "修改", "Repository 增加 findIdempotencyKey / saveIdempotencyKey"],
            ["shared/repository-memory.ts", "修改", "内存实现幂等键存取"],
            ["shared/services/profile-service.ts", "修改", "移除姓名去重；createProfile 支持 requestId 幂等；ClientProfile 增 isDefault；toClientProfile 计算 isDefault"],
            ["cloudfunctions/profileApi/cloudbase-repository.js", "修改", "idempotency_keys 集合的查/写实现"],
            ["cloudfunctions/profileApi/index.js", "修改", "currentDefaultId 助手；各分支传入 defaultId；create 透传 requestId"],
            ["cloudfunctions/login/cloudbase-repository.js", "修改", "接口对齐补齐幂等键方法"],
            ["typings/index.d.ts", "修改", "ClientFamilyProfile 增 isDefault（服务端计算）"],
            ["miniprogram/services/profile.ts", "修改", "newRequestId()；createProfile 传 requestId"],
            ["miniprogram/pages/profile-edit/profile-edit.ts", "修改", "onLoad 生成 createRequestId；onSubmit 透传"],
            ["scripts/validate.mjs", "修改", "以 14a/14b/14c + 15a–15f 替换原姓名去重测试；打包/卫生测试重编号 16a–d / 17"],
            ["docs/DATA_MODEL.md", "修改", "idempotency_keys、默认 SSoT、移除姓名去重"],
            ["docs/ARCHITECTURE.md", "修改", "§5b 幂等与默认 SSoT、§12 M1.1 决策、集合/索引表"],
            ["docs/USER_FLOWS.md", "修改", "创建流程 requestId、允许重名、默认计算与回退"],
            ["docs/SECURITY.md", "修改", "集合/索引、创建幂等与默认 SSoT 说明"],
            ["build_m1_report.py / 报告", "修改", "并入 M1.1 修正、测试状态与 git 复核"],
        ]),
        ("h2", "16.4 合并就绪结论"),
        ("p", "两处设计修正均已完成并通过自动化验收（91/91）。分支 feature/m1-identity-family-profiles 的**代码与文档已就绪，可进入产品负责人的 DevTools / 真机手动验收**（见第 17、18 节）。未宣称产品已验收通过——手动/真机冒烟（无法在当前自动化环境代跑）完成前，不应判定为 fully product-accepted。"),
    ]),
    ("17. 测试状态（自动化 / DevTools 手动 / 真机）", [
        ("p", "按要求严格区分三类测试的执行状态，避免将未执行项报告为已通过："),
        ("table", ["类别", "状态", "说明"], [
            ["自动化测试（npm run validate）", "✅ 已执行并通过", "91/91；含 M1.1 的 14a–c、15a–f"],
            ["DevTools 手动测试", "⏳ 待执行（产品负责人）", "需微信开发者工具环境，自动化环境无法代跑"],
            ["真机测试", "⏳ 待执行（产品负责人）", "需真机 + 已配置的 CloudBase dev 环境"],
        ]),
        ("h2", "17.1 DevTools 手动测试步骤（待执行）"),
        ("bullet", "1) 在 miniprogram/config/env.local.ts 填入 dev CloudBase 环境 ID；运行 npm run build:shared 生成各云函数 lib/shared/。"),
        ("bullet", "2) 用微信开发者工具打开项目，上传并部署 login、profileApi 两个云函数。"),
        ("bullet", "3) 首次运行：应进入引导；创建首个档案 → 自动成为默认。"),
        ("bullet", "4) 再创建第二个档案；在列表切换默认 → 仅一个显示为默认。"),
        ("bullet", "5) 编辑档案姓名/关系并保存；选择活动档案；重新编译，确认默认仍生效。"),
        ("bullet", "6) 提交空姓名/非法关系 → 被拒。"),
        ("bullet", "7) 【M1.1】连续快速点击提交 → 仅创建一条（requestId 幂等 + 防抖）。"),
        ("bullet", "8) 【M1.1】故意创建两条重名档案（各自独立提交）→ 均应成功、并列存在。"),
        ("h2", "17.2 真机测试步骤（待执行）"),
        ("bullet", "1) 在真机预览/体验版中复现 17.1 的 3–8。"),
        ("bullet", "2) 断网/未配置云环境时，首页应显示离线横幅且不崩溃。"),
        ("bullet", "3) 清除本地存储后重启，确认活动档案回退（服务端默认 → 首个 → 引导）正确。"),
        ("p", "以上步骤执行后，请将结果回填至 docs/MANUAL_TEST_RESULTS.md（手动测试结果表，14 项全 Pending）并在本节标注实际状态。"),
    ]),
    ("18. M1 最终 closeout（本交付修正）", [
        ("p", "本 closeout 是合入 main 前的最后一次窄范围修正，仅处理 PM 复核发现的报告与文档问题，不改变 M1/M1.1 的实现代码（实现已通过 91/91 自动化验收）。"),
        ("h2", "18.1 本轮修正内容"),
        ("bullet", "幂等表述：将「请求级幂等」统一改为 best-effort 请求幂等（尽力式请求幂等）/ 客户端 in-flight 防护 + 服务端请求重放处理；明确不宣称严格/原子幂等；在 DATA_MODEL / ARCHITECTURE / SECURITY / USER_FLOWS 与本报告记录残留并发竞态与 M3 通用原子方案。"),
        ("bullet", "Git 日志复核：第 14 节改为只列真实 hash 的 10 个提交，移除「(本次)」占位；补注本报告自身提交未计入日志（因其 hash 在生成前尚不存在）。"),
        ("bullet", "里程碑路线图：第 15 节改为约定序列 M2 食物目录与份量单位 → M3 手动记餐 → M4 历史/编辑/删除 → M5 收藏与食谱 → M6 照片 → M7 mock AI → M8 真实 AI；M2 明确排除 mealApi/餐史/食谱/照片/AI。同步更新 README 与 DEVELOPMENT_PLAN。"),
        ("bullet", "文件计数一致性：第 3 节改为「M1 最初影响 47 个文件（31 新增/16 修改）；M1.1 追加修改 18 个；本 closeout 修改一组文档/报告文件」，并分 §3.5 / §3.6 区分三类文件。"),
        ("bullet", "手动测试真实状态：新增 docs/MANUAL_TEST_RESULTS.md（14 项全 Pending）；第 17 节维持 DevTools / 真机为「待执行」。"),
        ("h2", "18.2 是否可合入 main"),
        ("p", "**代码与文档已就绪，可进入产品负责人的 DevTools / 真机手动验收。** 两处设计修正（移除姓名去重、默认档案单一事实来源）均已完成并通过自动化验收（91/91），含幂等与默认 SSoT 的关键项。唯一前置为产品负责人按 docs/MANUAL_TEST_RESULTS.md 执行手动/真机冒烟——该项受环境所限无法由自动化代跑，明确标注为 Pending，不构成代码合并阻塞，但在完成前不应判定为 fully product-accepted。本任务**未将分支合入 main**，也未创建 M2 分支。"),
    ]),
]


# ---------------------------------------------------------------------------
# Styling helpers
# ---------------------------------------------------------------------------
def set_run_font(run, name=FONT_NAME, size=FONT_SIZE, bold=False, color=None, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10.5, bold=True, color=WHITE)
        set_cell_bg(hdr_cells[i], '2B5C8F')
    for r_idx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=10)
            if r_idx % 2 == 1:
                set_cell_bg(cells[i], 'EAF1F8')
    # column widths: reasonable defaults
    widths = _col_widths(headers)
    for row in t.rows:
        for i, c in enumerate(row.cells):
            c.width = widths[i]
    return t


def _col_widths(headers):
    total = len(headers)
    base = Inches(6.5 / total)
    return [base] * total


# ---------------------------------------------------------------------------
# DOCX builder
# ---------------------------------------------------------------------------
def build_docx(path):
    doc = Document()
    # default style: 宋体 12pt, 1.5 line spacing
    normal = doc.styles['Normal']
    normal.font.name = FONT_NAME
    normal.font.size = Pt(FONT_SIZE)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), FONT_NAME)
    rfonts.set(qn('w:ascii'), FONT_NAME)
    rfonts.set(qn('w:hAnsi'), FONT_NAME)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5

    # ---- cover page ----
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    set_run_font(run, size=28, bold=True, color=BLUE_DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(SUBTITLE)
    set_run_font(run, size=16, bold=True, color=BLUE_HEAD)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("最终交付报告")
    set_run_font(run, size=14, color=BLUE_HEAD)
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"日期：{REPORT_DATE}")
    set_run_font(run, size=12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("状态：✅ 代码与文档完成（91/91 自动化校验通过；DevTools/真机手动验收待产品负责人）")
    set_run_font(run, size=12, bold=True, color=BLUE_DARK)
    doc.add_page_break()

    # ---- table of contents (manual) ----
    p = doc.add_paragraph()
    run = p.add_run("目录")
    set_run_font(run, size=18, bold=True, color=BLUE_DARK)
    for idx, (title, _) in enumerate(SECTIONS, start=1):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(f"{idx}. {title.split('.', 1)[1].strip()}")
        set_run_font(run, size=12)
    doc.add_page_break()

    # ---- sections ----
    for title, blocks in SECTIONS:
        # section heading
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(6)
        run = h.add_run(title)
        set_run_font(run, size=15, bold=True, color=BLUE_DARK)
        for block in blocks:
            kind = block[0]
            payload = block[1] if len(block) < 3 else None
            if kind == 'table':
                _, headers, rows = block
                add_table(doc, headers, rows)
                doc.add_paragraph()
                continue
            if kind == 'h2':
                ph = doc.add_paragraph()
                ph.paragraph_format.space_before = Pt(6)
                ph.paragraph_format.space_after = Pt(3)
                run = ph.add_run(payload)
                set_run_font(run, size=12.5, bold=True, color=BLUE_HEAD)
            elif kind == 'p':
                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(6)
                run = para.add_run(payload)
                set_run_font(run, size=12)
            elif kind == 'bullet':
                para = doc.add_paragraph(style='List Bullet')
                para.paragraph_format.space_after = Pt(3)
                run = para.add_run(payload)
                set_run_font(run, size=12)
            elif kind == 'table':
                headers, rows = payload
                add_table(doc, headers, rows)
                doc.add_paragraph()
            elif kind == 'code':
                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(6)
                run = para.add_run(payload)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                # light shading via paragraph border-ish: use a single-cell table for bg
                _code_block(doc, payload)

    doc.save(path)


def _code_block(doc, text):
    # render code in a shaded single-cell table for readability
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, 'F2F4F7')
    cell.text = ''
    p = cell.paragraphs[0]
    for ln in text.split('\n'):
        r = p.add_run(ln + '\n')
        r.font.name = 'Consolas'
        r.font.size = Pt(8.5)
        rpr = r._element.get_or_add_rPr()
        rf = rpr.find(qn('w:rFonts'))
        if rf is None:
            rf = OxmlElement('w:rFonts')
            rpr.append(rf)
        rf.set(qn('w:ascii'), 'Consolas')
        rf.set(qn('w:hAnsi'), 'Consolas')
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Markdown builder
# ---------------------------------------------------------------------------
def build_md(path):
    out = []
    out.append(f"# {TITLE} — {SUBTITLE}（最终交付报告）\n")
    out.append(f"> 日期：{REPORT_DATE} ｜ 状态：✅ 代码与文档完成（91/91 自动化校验通过；DevTools/真机手动验收待产品负责人）\n")
    out.append("\n## 目录\n")
    for idx, (title, _) in enumerate(SECTIONS, start=1):
        out.append(f"{idx}. {title.split('.', 1)[1].strip()}")
    out.append("")
    for title, blocks in SECTIONS:
        out.append(f"## {title}\n")
        for block in blocks:
            kind = block[0]
            payload = block[1] if len(block) < 3 else None
            if kind == 'h2':
                out.append(f"### {payload}\n")
            elif kind == 'p':
                out.append(payload + "\n")
            elif kind == 'bullet':
                out.append(f"- {payload}")
            elif kind == 'table':
                _, headers, rows = block
                out.append("| " + " | ".join(headers) + " |")
                out.append("| " + " | ".join(["---"] * len(headers)) + " |")
                for row in rows:
                    out.append("| " + " | ".join(str(c) for c in row) + " |")
                out.append("")
            elif kind == 'code':
                out.append("```")
                out.append(payload)
                out.append("```\n")
    with open(path, 'w', encoding='utf-8') as f:
        f.write("\n".join(out))


if __name__ == '__main__':
    base = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(base, 'docs', 'M1-FINAL-REPORT.md')
    docx_path = os.path.join(base, 'M1-FINAL-REPORT.docx')
    os.makedirs(os.path.dirname(md_path), exist_ok=True)
    build_md(md_path)
    build_docx(docx_path)
    print("WROTE", md_path)
    print("WROTE", docx_path)
